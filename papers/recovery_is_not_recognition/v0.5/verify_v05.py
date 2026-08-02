#!/usr/bin/env python3
from pathlib import Path
import json
from docx import Document
import fitz

B=Path(__file__).resolve().parent
checks=[]
def check(name,cond,detail=''):
    checks.append({'name':name,'pass':bool(cond),'detail':str(detail)})

files=[
'manuscript_v0_5.md','cryptologia_recovery_not_recognition_v0_5.docx',
'cryptologia_recovery_not_recognition_v0_5.tex','cryptologia_recovery_not_recognition_v0_5.pdf',
'cryptologia_recovery_not_recognition_v0_5_word.pdf','README.md','CHANGELOG_v0_5.md',
'PAPER_REVISION_PROTOCOL_v0_5.md','CLAIM_LEDGER_v0_5.md','REFERENCE_AUDIT_v0_5.md',
'COLD_REVIEW_v0_5.md','BUILD_REPORT_v0_5.md','MANIFEST_v0_5.md','V0_5_REPAIR_AUDIT.md',
'a11y_report_v0_5.json','style_lint_v0_5.txt','heading_audit_v0_5.txt','pdf_preflight_v0_5.txt',
'SOURCE_HASH_RECORD_v0_5.md']
for f in files: check('file:'+f,(B/f).is_file() and (B/f).stat().st_size>0,(B/f).stat().st_size if (B/f).exists() else 0)

s=(B/'manuscript_v0_5.md').read_text()
required=[
'From surface compatibility to evidential calibration',
'Positional structure and the operational four-part representation',
'Adversarial falsification of the aggregate score',
'Held-out two-sample discrimination',
'A focused within-line order test',
'What generator success can and cannot establish',
'Transliteration as a measurement model',
'Constructive compatibility is not historical identification',
'Generator claims require untouched discrimination',
'line-shuffled control', '74.6 ± 0.7', '66.9 ± 0.3',
'AUC 0.485 ± 0.063','0.872 ± 0.053','0.992 ± 0.008',
'All eleven natural texts','best seed reaching 0.87',
'not optimised','one operational exemplar','bounded existence result',
'A fitted generator supplies a bounded existence proof',
'These results do not classify the Voynich Manuscript'
]
for x in required: check('text:'+x,x in s)
check('no +/- typography','+/-' not in s)
check('no raw tool citations','filecite' not in s and 'turn3' not in s)
check('references count',len([x for x in s.split('# References',1)[1].strip().split('\n\n') if x.strip()])==38)
check('author anonymous','author: "Anonymous manuscript for review"' in s)
check('H2 archive excluded from manifest','vms_h2_archive_2026-07-12.zip' not in (B/'MANIFEST_v0_5.md').read_text())

D=Document(B/'cryptologia_recovery_not_recognition_v0_5.docx')
check('docx paragraphs',len(D.paragraphs)>300,len(D.paragraphs))
check('docx tables',len(D.tables)==1,len(D.tables))
check('docx table rows',len(D.tables[0].rows)==19,len(D.tables[0].rows))

for f,expected in [('cryptologia_recovery_not_recognition_v0_5.pdf',32),('cryptologia_recovery_not_recognition_v0_5_word.pdf',63)]:
    d=fitz.open(B/f)
    check(f+' pages',d.page_count==expected,d.page_count)
    text=''.join(p.get_text() for p in d)
    check(f+' text-based',len(text)>80000,len(text))

aj=json.loads((B/'a11y_report_v0_5.json').read_text())
raw=json.dumps(aj).lower()
check('a11y no high','"high": 0' in raw or '"high_count": 0' in raw,raw[:250])
check('a11y no medium','"medium": 0' in raw or '"medium_count": 0' in raw,raw[:250])

shr=(B/'SOURCE_HASH_RECORD_v0_5.md').read_text()
check('H2 hash recorded','8e7f6205154990db88b19fe3c378fcabb43a55a5f056d1e2897370aff2062a39' in shr)
check('demolition hash recorded','588479993113632fb171844ba4141991011dff4ae20583e7744416cdb761ace2' in shr)

passed=sum(c['pass'] for c in checks)
result={'status':'PASS' if passed==len(checks) else 'FAIL','passed':passed,'total':len(checks),'checks':checks}
(B/'VERIFY_RESULT_v0_5.json').write_text(json.dumps(result,indent=2,ensure_ascii=False)+'\n')
print(json.dumps({'status':result['status'],'passed':passed,'total':len(checks)},indent=2))
if result['status']!='PASS':
    for c in checks:
        if not c['pass']: print('FAIL',c)
    raise SystemExit(1)
